"""
Legal Document Analyzer - In-Memory Processing Only
Analyzes uploaded legal documents (wills, agreements, contracts) without storing data
"""

import os
import sys
import io
import re
from typing import Dict, List, Optional, Tuple
import PyPDF2
import docx
from datetime import datetime, timedelta

# Add ml_legal_system to path
sys.path.append(os.path.join(os.path.dirname(__file__), 'ml_legal_system'))

try:
    from ml_legal_system.legal_rag import LegalRAG
    ML_AVAILABLE = True
except ImportError:
    ML_AVAILABLE = False


class DocumentAnalyzer:
    """
    Analyzes legal documents in-memory without persistent storage
    Supports: PDF, DOCX, TXT
    """
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB limit
    
    def __init__(self):
        """Initialize document analyzer"""
        self.rag = None
        if ML_AVAILABLE:
            try:
                self.rag = LegalRAG(use_openai=False)
            except Exception as e:
                print(f"⚠️  Could not initialize RAG: {e}")
    
    def validate_file(self, filename: str, file_size: int) -> Dict:
        """
        Validate uploaded file
        
        Args:
            filename: Name of the file
            file_size: Size of file in bytes
            
        Returns:
            Dict with 'valid' boolean and 'error' message if invalid
        """
        # Check file extension
        file_ext = os.path.splitext(filename.lower())[1]
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return {
                'valid': False,
                'error': f'Unsupported file type. Please upload PDF, DOCX, or TXT files.'
            }
        
        # Check file size
        if file_size > self.MAX_FILE_SIZE:
            return {
                'valid': False,
                'error': f'File too large. Maximum size is 10 MB.'
            }
        
        return {'valid': True}
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            
            return '\n\n'.join(text)
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n\n'.join(text)
            
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return file_content.decode('latin-1')
            except Exception as e:
                raise ValueError(f"Failed to decode text file: {str(e)}")
    
    def extract_text(self, filename: str, file_content: bytes) -> str:
        """
        Extract text from uploaded file based on extension
        
        Args:
            filename: Name of the file
            file_content: Raw file bytes
            
        Returns:
            Extracted text content
        """
        file_ext = os.path.splitext(filename.lower())[1]
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_content)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def detect_document_type(self, text: str) -> str:
        """
        Detect type of legal document based on content
        
        Returns:
            Document type (will, agreement, contract, etc.)
        """
        text_lower = text.lower()
        
        # Keywords for different document types
        if any(word in text_lower for word in ['last will', 'testament', 'testator', 'executor', 'bequeath']):
            return 'Will / Testament'
        elif any(word in text_lower for word in ['rental agreement', 'lease agreement', 'tenant', 'landlord', 'rent']):
            return 'Rental Agreement'
        elif any(word in text_lower for word in ['employment agreement', 'employment contract', 'employee', 'employer', 'salary']):
            return 'Employment Agreement'
        elif any(word in text_lower for word in ['sale agreement', 'purchase agreement', 'buyer', 'seller', 'consideration']):
            return 'Sale Agreement'
        elif any(word in text_lower for word in ['partnership agreement', 'partners', 'partnership deed']):
            return 'Partnership Agreement'
        elif any(word in text_lower for word in ['loan agreement', 'borrower', 'lender', 'principal amount']):
            return 'Loan Agreement'
        elif any(word in text_lower for word in ['non-disclosure agreement', 'nda', 'confidential information']):
            return 'Non-Disclosure Agreement (NDA)'
        elif any(word in text_lower for word in ['service agreement', 'service provider', 'client']):
            return 'Service Agreement'
        elif 'agreement' in text_lower or 'contract' in text_lower:
            return 'General Agreement/Contract'
        else:
            return 'Legal Document'
    
    def extract_key_information(self, text: str, doc_type: str) -> Dict:
        """
        Extract key information from document based on type
        
        Returns:
            Dictionary with extracted key information
        """
        info = {
            'document_type': doc_type,
            'parties': [],
            'dates': [],
            'key_terms': []
        }
        
        # Extract dates (simple pattern matching)
        import re
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # DD/MM/YYYY or DD-MM-YYYY
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY-MM-DD
            r'\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',  # DD Month YYYY
        ]
        
        for pattern in date_patterns:
            dates = re.findall(pattern, text, re.IGNORECASE)
            info['dates'].extend(dates[:3])  # Limit to first 3 dates
        
        # Extract party names (simplified - looks for common patterns)
        party_keywords = ['party of the first part', 'party of the second part', 'between', 'and']
        lines = text.split('\n')
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in party_keywords):
                # Check next few lines for potential party names
                for j in range(i, min(i+3, len(lines))):
                    if lines[j].strip() and len(lines[j].strip().split()) <= 5:
                        potential_party = lines[j].strip()
                        if potential_party and not any(kw in potential_party.lower() for kw in party_keywords):
                            info['parties'].append(potential_party)
        
        # Extract key terms based on document type
        if 'will' in doc_type.lower():
            info['key_terms'] = ['Testator', 'Executor', 'Beneficiaries', 'Bequests', 'Distribution of Assets']
        elif 'rental' in doc_type.lower() or 'lease' in doc_type.lower():
            info['key_terms'] = ['Rent Amount', 'Security Deposit', 'Lease Duration', 'Notice Period', 'Maintenance']
        elif 'employment' in doc_type.lower():
            info['key_terms'] = ['Salary/Compensation', 'Job Title', 'Notice Period', 'Confidentiality', 'Non-compete']
        elif 'sale' in doc_type.lower():
            info['key_terms'] = ['Sale Price', 'Payment Terms', 'Delivery Date', 'Warranties', 'Possession']
        else:
            info['key_terms'] = ['Obligations', 'Payment Terms', 'Termination', 'Liability', 'Dispute Resolution']
        
        return info
    
    def extract_contract_clauses(self, text: str, doc_type: str) -> Dict:
        """
        Extract and analyze key contract clauses from legal documents
        
        Returns:
            Dictionary with identified clauses and their content
        """
        # Define legal clause patterns and keywords
        clause_patterns = {
            'termination': {
                'keywords': ['termination', 'terminate', 'end of agreement', 'expiry', 'dissolution'],
                'context_phrases': ['may be terminated', 'shall terminate', 'termination notice', 'upon termination'],
                'importance': 'high'
            },
            'payment': {
                'keywords': ['payment', 'amount', 'consideration', 'fee', 'salary', 'compensation'],
                'context_phrases': ['shall pay', 'payment due', 'amount of', 'consideration for'],
                'importance': 'high'
            },
            'liability': {
                'keywords': ['liability', 'liable', 'damages', 'responsible', 'indemnify', 'indemnification'],
                'context_phrases': ['shall be liable', 'liability for', 'damages arising', 'indemnify against'],
                'importance': 'high'
            },
            'confidentiality': {
                'keywords': ['confidential', 'non-disclosure', 'proprietary', 'trade secret'],
                'context_phrases': ['confidential information', 'shall not disclose', 'proprietary information'],
                'importance': 'medium'
            },
            'jurisdiction': {
                'keywords': ['jurisdiction', 'governing law', 'court', 'dispute resolution', 'arbitration'],
                'context_phrases': ['governed by', 'jurisdiction of', 'subject to the laws'],
                'importance': 'medium'
            },
            'force_majeure': {
                'keywords': ['force majeure', 'act of god', 'unforeseen circumstances', 'beyond control'],
                'context_phrases': ['force majeure', 'acts of god', 'circumstances beyond'],
                'importance': 'medium'
            },
            'intellectual_property': {
                'keywords': ['intellectual property', 'copyright', 'trademark', 'patent', 'proprietary rights'],
                'context_phrases': ['intellectual property rights', 'proprietary rights', 'copyright in'],
                'importance': 'medium'
            },
            'warranty': {
                'keywords': ['warranty', 'warrants', 'guarantee', 'representation'],
                'context_phrases': ['warrants that', 'represents and warrants', 'guarantee that'],
                'importance': 'medium'
            }
        }
        
        identified_clauses = {}
        text_lower = text.lower()
        sentences = self._split_into_sentences(text)
        
        for clause_type, patterns in clause_patterns.items():
            clause_content = []
            
            # Look for sentences containing clause keywords or context phrases
            for sentence in sentences:
                sentence_lower = sentence.lower()
                
                # Check for direct keyword matches
                keyword_matches = any(keyword in sentence_lower for keyword in patterns['keywords'])
                context_matches = any(phrase in sentence_lower for phrase in patterns['context_phrases'])
                
                if keyword_matches or context_matches:
                    # Clean and add the sentence
                    clean_sentence = sentence.strip()
                    if len(clean_sentence) > 20:  # Filter out very short matches
                        clause_content.append(clean_sentence)
            
            if clause_content:
                identified_clauses[clause_type] = {
                    'content': clause_content[:3],  # Limit to top 3 most relevant sentences
                    'importance': patterns['importance'],
                    'count': len(clause_content)
                }
        
        return identified_clauses
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for clause analysis"""
        # Simple sentence splitting - can be enhanced with NLTK
        sentence_endings = r'[.!?]+\s+'
        sentences = re.split(sentence_endings, text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            clean = sentence.strip()
            if len(clean) > 10:  # Filter very short fragments
                cleaned_sentences.append(clean)
        
        return cleaned_sentences
    
    def extract_key_dates_and_deadlines(self, text: str) -> Dict:
        """
        Extract important dates and calculate deadlines from legal documents
        
        Returns:
            Dictionary with dates, deadlines, and time-sensitive information
        """
        date_info = {
            'identified_dates': [],
            'deadlines': [],
            'time_sensitive_clauses': [],
            'notice_periods': []
        }
        
        # Enhanced date pattern matching
        date_patterns = [
            (r'\b(\d{1,2})[-/](\d{1,2})[-/](\d{2,4})\b', 'DD/MM/YYYY'),
            (r'\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b', 'YYYY-MM-DD'),
            (r'\b(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{4})\b', 'DD Month YYYY'),
            (r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b', 'Month DD, YYYY'),
        ]
        
        # Find all dates
        for pattern, format_type in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                context_start = max(0, match.start() - 50)
                context_end = min(len(text), match.end() + 50)
                context = text[context_start:context_end].strip()
                
                date_info['identified_dates'].append({
                    'date': match.group(),
                    'format': format_type,
                    'context': context
                })
        
        # Look for deadline-related phrases
        deadline_patterns = [
            r'within\s+(\d+)\s+(days?|weeks?|months?|years?)',
            r'not later than\s+(\d+)\s+(days?|weeks?|months?|years?)',
            r'deadline of\s+(\d+)\s+(days?|weeks?|months?|years?)',
            r'expires?\s+(?:on|in)\s+(\d+)\s+(days?|weeks?|months?|years?)',
        ]
        
        for pattern in deadline_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                context_start = max(0, match.start() - 100)
                context_end = min(len(text), match.end() + 100)
                context = text[context_start:context_end].strip()
                
                date_info['deadlines'].append({
                    'period': match.group(),
                    'context': context
                })
        
        # Look for notice periods
        notice_patterns = [
            r'(\d+)\s+(days?|weeks?|months?)\s+notice',
            r'notice\s+of\s+(\d+)\s+(days?|weeks?|months?)',
            r'(\d+)\s+(days?|weeks?|months?)\s+prior\s+notice',
        ]
        
        for pattern in notice_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                context_start = max(0, match.start() - 80)
                context_end = min(len(text), match.end() + 80)
                context = text[context_start:context_end].strip()
                
                date_info['notice_periods'].append({
                    'period': match.group(),
                    'context': context
                })
        
        return date_info
    
    def assess_legal_risks(self, text: str, doc_type: str, clauses: Dict) -> Dict:
        """
        Assess potential legal risks in the document
        
        Returns:
            Dictionary with risk assessment and recommendations
        """
        risks = {
            'high_risk_issues': [],
            'medium_risk_issues': [],
            'low_risk_issues': [],
            'missing_protections': [],
            'overall_risk_score': 0
        }
        
        text_lower = text.lower()
        
        # High-risk indicators
        high_risk_patterns = [
            ('unlimited liability', 'Document may contain unlimited liability clauses'),
            ('waiver of rights', 'Rights waiver clauses detected'),
            ('no warranty', 'Warranty disclaimers present'),
            ('sole discretion', 'One-sided discretionary powers identified'),
            ('non-refundable', 'Non-refundable payment terms'),
        ]
        
        # Medium-risk indicators  
        medium_risk_patterns = [
            ('automatic renewal', 'Automatic renewal clauses present'),
            ('penalty', 'Penalty clauses identified'),
            ('liquidated damages', 'Liquidated damages provisions'),
            ('assignment', 'Assignment rights may be unclear'),
            ('modification', 'Document modification terms'),
        ]
        
        # Low-risk indicators
        low_risk_patterns = [
            ('subject to approval', 'Approval-dependent clauses'),
            ('reasonable efforts', 'Performance standards may be vague'),
            ('best efforts', 'Effort-based obligations'),
        ]
        
        # Check for risk patterns
        for pattern, description in high_risk_patterns:
            if pattern in text_lower:
                risks['high_risk_issues'].append(description)
        
        for pattern, description in medium_risk_patterns:
            if pattern in text_lower:
                risks['medium_risk_issues'].append(description)
        
        for pattern, description in low_risk_patterns:
            if pattern in text_lower:
                risks['low_risk_issues'].append(description)
        
        # Check for missing critical clauses
        if 'termination' not in clauses:
            risks['missing_protections'].append('No clear termination clause identified')
        
        if 'liability' not in clauses:
            risks['missing_protections'].append('Liability limitations not found')
        
        if 'jurisdiction' not in clauses:
            risks['missing_protections'].append('Governing law and jurisdiction not specified')
        
        if doc_type.lower() in ['employment', 'service'] and 'confidentiality' not in clauses:
            risks['missing_protections'].append('Confidentiality provisions missing')
        
        # Calculate overall risk score (0-100)
        risk_score = 0
        risk_score += len(risks['high_risk_issues']) * 20
        risk_score += len(risks['medium_risk_issues']) * 10
        risk_score += len(risks['low_risk_issues']) * 5
        risk_score += len(risks['missing_protections']) * 15
        
        risks['overall_risk_score'] = min(100, risk_score)
        
        return risks

    def analyze_document(self, filename: str, file_content: bytes,
                        specific_questions: Optional[List[str]] = None) -> Dict:
        """
        Analyze legal document and provide insights
        
        Args:
            filename: Name of the uploaded file
            file_content: Raw file bytes
            specific_questions: Optional list of specific questions to answer
            
        Returns:
            Dictionary with analysis results (NOT stored anywhere)
        """
        try:
            # Validate file
            validation = self.validate_file(filename, len(file_content))
            if not validation['valid']:
                return {
                    'success': False,
                    'error': validation['error']
                }
            
            # Extract text
            text = self.extract_text(filename, file_content)
            
            if not text or len(text.strip()) < 50:
                return {
                    'success': False,
                    'error': 'Could not extract sufficient text from document. File may be empty or corrupted.'
                }
            
            # Detect document type
            doc_type = self.detect_document_type(text)
            
            # Extract key information
            key_info = self.extract_key_information(text, doc_type)
            
            # ✨ NEW: Extract contract clauses
            contract_clauses = self.extract_contract_clauses(text, doc_type)
            
            # ✨ NEW: Extract dates and deadlines
            date_analysis = self.extract_key_dates_and_deadlines(text)
            
            # ✨ NEW: Assess legal risks
            risk_assessment = self.assess_legal_risks(text, doc_type, contract_clauses)
            
            # Generate analysis using RAG if available
            analysis = self._generate_analysis(text, doc_type, specific_questions)
            
            # ✨ NEW: Generate document summary
            document_summary = self._generate_document_summary(text, doc_type, contract_clauses)
            
            # Calculate statistics
            word_count = len(text.split())
            char_count = len(text)
            
            # Return enhanced analysis (in-memory only, NOT stored)
            return {
                'success': True,
                'filename': filename,
                'document_type': doc_type,
                'statistics': {
                    'word_count': word_count,
                    'character_count': char_count,
                    'pages_estimated': max(1, word_count // 300)
                },
                'key_information': key_info,
                'contract_clauses': contract_clauses,
                'date_analysis': date_analysis,
                'risk_assessment': risk_assessment,
                'document_summary': document_summary,
                'analysis': analysis,
                'timestamp': datetime.now().isoformat(),
                'disclaimer': 'This analysis is for informational purposes only and does not constitute legal advice. Consult a qualified attorney for specific guidance.'
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Document analysis failed: {str(e)}'
            }
    
    def _generate_analysis(self, text: str, doc_type: str, 
                          specific_questions: Optional[List[str]] = None) -> Dict:
        """
        Generate detailed analysis of document using RAG
        
        Returns:
            Dictionary with analysis insights
        """
        analysis = {
            'summary': '',
            'key_clauses': [],
            'potential_issues': [],
            'recommendations': [],
            'answers_to_questions': []
        }
        
        # Generate summary
        if self.rag:
            try:
                summary_query = f"Provide a brief summary of this {doc_type}: {text[:1500]}"
                summary_result = self.rag.answer_legal_query(summary_query, top_k=3)
                analysis['summary'] = summary_result['answer']
            except Exception as e:
                print(f"⚠️  RAG summary failed: {e}")
                analysis['summary'] = self._generate_basic_summary(text, doc_type)
        else:
            analysis['summary'] = self._generate_basic_summary(text, doc_type)
        
        # Identify key clauses based on document type
        analysis['key_clauses'] = self._identify_key_clauses(text, doc_type)
        
        # Identify potential issues
        analysis['potential_issues'] = self._identify_potential_issues(text, doc_type)
        
        # Generate recommendations
        analysis['recommendations'] = self._generate_recommendations(doc_type)
        
        # Answer specific questions if provided
        if specific_questions and self.rag:
            for question in specific_questions[:5]:  # Limit to 5 questions
                try:
                    # Context-aware query with document text
                    query = f"Based on this document: {text[:2000]}... {question}"
                    result = self.rag.answer_legal_query(query, top_k=3)
                    analysis['answers_to_questions'].append({
                        'question': question,
                        'answer': result['answer']
                    })
                except Exception as e:
                    print(f"⚠️  Question answering failed: {e}")
        
        return analysis
    
    def _generate_document_summary(self, text: str, doc_type: str, clauses: Dict) -> Dict:
        """
        Generate a comprehensive document summary with key highlights
        
        Returns:
            Dictionary with document summary and key insights
        """
        summary = {
            'overview': '',
            'key_highlights': [],
            'important_clauses': [],
            'action_items': [],
            'compliance_notes': []
        }
        
        # Generate overview based on document type
        if 'employment' in doc_type.lower():
            summary['overview'] = f"Employment agreement document with {len(text.split())} words. Contains employment terms, conditions, and obligations for both employer and employee."
        elif 'rental' in doc_type.lower() or 'lease' in doc_type.lower():
            summary['overview'] = f"Rental/lease agreement document outlining terms between landlord and tenant for property rental arrangements."
        elif 'sale' in doc_type.lower():
            summary['overview'] = f"Sale agreement document establishing terms for transfer of ownership between buyer and seller."
        elif 'will' in doc_type.lower():
            summary['overview'] = f"Last Will and Testament document outlining distribution of assets and final wishes of the testator."
        else:
            summary['overview'] = f"{doc_type} document containing legal terms, conditions, and obligations between parties."
        
        # Extract key highlights from important clauses
        important_clause_types = ['termination', 'payment', 'liability']
        for clause_type in important_clause_types:
            if clause_type in clauses and clauses[clause_type]['content']:
                summary['important_clauses'].append({
                    'type': clause_type.replace('_', ' ').title(),
                    'content': clauses[clause_type]['content'][0][:200] + "..."
                })
        
        # Generate key highlights
        text_lower = text.lower()
        
        # Look for monetary amounts
        money_pattern = r'\$[\d,]+(?:\.\d{2})?|\b(?:rupees?|rs\.?)\s*[\d,]+(?:\.\d{2})?'
        money_matches = re.findall(money_pattern, text, re.IGNORECASE)
        if money_matches:
            summary['key_highlights'].append(f"Financial terms: {', '.join(money_matches[:3])}")
        
        # Look for time periods
        time_pattern = r'\b\d+\s+(?:days?|weeks?|months?|years?)\b'
        time_matches = re.findall(time_pattern, text, re.IGNORECASE)
        if time_matches:
            summary['key_highlights'].append(f"Time periods: {', '.join(set(time_matches[:3]))}")
        
        # Generate action items based on document type
        if 'employment' in doc_type.lower():
            summary['action_items'] = [
                "Review compensation and benefits terms",
                "Understand notice period requirements", 
                "Check confidentiality and non-compete clauses",
                "Verify job responsibilities and reporting structure"
            ]
        elif 'rental' in doc_type.lower():
            summary['action_items'] = [
                "Verify rent amount and due dates",
                "Review security deposit terms",
                "Understand maintenance responsibilities",
                "Check notice period for termination"
            ]
        elif 'sale' in doc_type.lower():
            summary['action_items'] = [
                "Verify purchase price and payment terms",
                "Check delivery and possession dates", 
                "Review warranty and return policies",
                "Understand transfer of ownership process"
            ]
        else:
            summary['action_items'] = [
                "Review all financial obligations",
                "Understand termination procedures",
                "Check dispute resolution mechanisms",
                "Verify all parties' responsibilities"
            ]
        
        # Add compliance notes
        summary['compliance_notes'] = [
            "Ensure all parties have legal capacity to enter the agreement",
            "Verify compliance with applicable local and federal laws",
            "Consider tax implications of the agreement terms",
            "Review insurance and liability coverage requirements"
        ]
        
        return summary
    
    def _generate_basic_summary(self, text: str, doc_type: str) -> str:
        """Generate basic summary without RAG"""
        return f"""This appears to be a {doc_type}. The document contains approximately {len(text.split())} words and covers legal terms and obligations between the parties involved. Key sections should be reviewed carefully, particularly regarding rights, obligations, payment terms, and dispute resolution mechanisms."""
    
    def _identify_key_clauses(self, text: str, doc_type: str) -> List[str]:
        """Identify important clauses in the document"""
        clauses = []
        text_lower = text.lower()
        
        # Common important clause indicators
        if 'termination' in text_lower or 'cancellation' in text_lower:
            clauses.append("Termination/Cancellation Clause - Review conditions for ending the agreement")
        
        if 'liability' in text_lower or 'indemnity' in text_lower:
            clauses.append("Liability/Indemnity Clause - Understand who is responsible for what")
        
        if 'dispute' in text_lower or 'arbitration' in text_lower or 'jurisdiction' in text_lower:
            clauses.append("Dispute Resolution Clause - Know how conflicts will be resolved")
        
        if 'confidential' in text_lower or 'non-disclosure' in text_lower:
            clauses.append("Confidentiality Clause - Understand privacy obligations")
        
        if 'payment' in text_lower or 'consideration' in text_lower or 'price' in text_lower:
            clauses.append("Payment Terms - Verify amounts, due dates, and payment methods")
        
        if 'warranty' in text_lower or 'guarantee' in text_lower:
            clauses.append("Warranty/Guarantee Clause - Check what is guaranteed and for how long")
        
        return clauses if clauses else ["Review all clauses carefully with legal counsel"]
    
    def _identify_potential_issues(self, text: str, doc_type: str) -> List[str]:
        """Identify potential issues or red flags"""
        issues = []
        text_lower = text.lower()
        
        # Check for missing important elements
        if 'agreement' in doc_type.lower() or 'contract' in doc_type.lower():
            if 'date' not in text_lower and 'dated' not in text_lower:
                issues.append("⚠️ Document may be missing execution date")
            
            if 'sign' not in text_lower and 'signature' not in text_lower:
                issues.append("⚠️ No signature section visible - verify document is properly executed")
        
        # Check for one-sided terms
        if 'non-refundable' in text_lower:
            issues.append("⚠️ Contains non-refundable terms - ensure you understand implications")
        
        if 'sole discretion' in text_lower:
            issues.append("⚠️ Contains 'sole discretion' clause - one party may have unilateral power")
        
        if 'perpetual' in text_lower or 'indefinite' in text_lower:
            issues.append("⚠️ May contain perpetual/indefinite obligations - check if there's an exit clause")
        
        return issues if issues else ["No obvious red flags detected - still recommend legal review"]
    
    def _generate_recommendations(self, doc_type: str) -> List[str]:
        """Generate recommendations based on document type"""
        recommendations = [
            "✓ Have this document reviewed by a qualified attorney before signing",
            "✓ Ensure all parties understand their rights and obligations",
            "✓ Keep a signed copy for your records"
        ]
        
        if 'will' in doc_type.lower():
            recommendations.extend([
                "✓ Ensure the will is properly witnessed and notarized",
                "✓ Review beneficiary designations regularly",
                "✓ Store the original in a safe place and inform executor of its location"
            ])
        elif 'rental' in doc_type.lower() or 'lease' in doc_type.lower():
            recommendations.extend([
                "✓ Inspect the property before signing",
                "✓ Document the property condition with photos",
                "✓ Understand the notice period for termination"
            ])
        elif 'employment' in doc_type.lower():
            recommendations.extend([
                "✓ Verify the job description and responsibilities",
                "✓ Understand non-compete and confidentiality obligations",
                "✓ Clarify benefits and termination procedures"
            ])
        
        return recommendations


# Singleton instance
_analyzer = None


def get_document_analyzer() -> DocumentAnalyzer:
    """Get or create document analyzer instance"""
    global _analyzer
    if _analyzer is None:
        _analyzer = DocumentAnalyzer()
    return _analyzer


if __name__ == "__main__":
    """Test document analyzer"""
    print("🔍 Document Analyzer - Test Mode")
    print("=" * 60)
    
    analyzer = get_document_analyzer()
    print(f"✅ Analyzer initialized")
    print(f"📋 Supported formats: {', '.join(analyzer.SUPPORTED_EXTENSIONS)}")
    print(f"📏 Max file size: {analyzer.MAX_FILE_SIZE / (1024*1024):.0f} MB")
    print("\n✅ Ready to analyze legal documents in-memory (no storage)")
