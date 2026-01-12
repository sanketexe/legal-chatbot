"""
Flask Blueprint for UX Enhancements (Task 5)
API endpoints for conversation history, bookmarks, and exports
"""

from flask import Blueprint, request, jsonify, send_file, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from models import db, ChatSession, Message, Bookmark, ExportHistory, User
from datetime import datetime
import traceback
import os
import json
from pathlib import Path

# Create blueprint
ux_bp = Blueprint('ux', __name__, url_prefix='/api/ux')


# ==================== CONVERSATION HISTORY ====================

@ux_bp.route('/conversations', methods=['GET'])
@jwt_required()
def get_conversations():
    """
    Get user's conversation history
    
    Query params:
    - limit: Number of conversations (default: 20)
    - offset: Pagination offset (default: 0)
    - active_only: Show only active conversations (default: true)
    """
    try:
        user_id = get_jwt_identity()
        limit = request.args.get('limit', 20, type=int)
        offset = request.args.get('offset', 0, type=int)
        active_only = request.args.get('active_only', 'true').lower() == 'true'
        
        # Build query
        query = ChatSession.query.filter_by(user_id=user_id)
        
        if active_only:
            query = query.filter_by(is_active=True)
        
        # Get total count
        total = query.count()
        
        # Get paginated results
        sessions = query.order_by(
            ChatSession.updated_at.desc()
        ).limit(limit).offset(offset).all()
        
        return jsonify({
            'success': True,
            'conversations': [s.to_dict() for s in sessions],
            'total': total,
            'limit': limit,
            'offset': offset
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


@ux_bp.route('/conversations/<session_id>', methods=['GET'])
@jwt_required()
def get_conversation(session_id):
    """Get specific conversation with all messages"""
    try:
        user_id = get_jwt_identity()
        
        session = ChatSession.query.filter_by(
            id=session_id,
            user_id=user_id
        ).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Conversation not found'
            }), 404
        
        return jsonify({
            'success': True,
            'conversation': session.to_dict(include_messages=True)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/conversations', methods=['POST'])
@jwt_required()
def create_conversation():
    """
    Create a new conversation
    
    Request body:
    {
        "title": "Optional title"
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json() or {}
        
        session = ChatSession(
            user_id=user_id,
            title=data.get('title')
        )
        
        db.session.add(session)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'conversation': session.to_dict()
        }), 201
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/conversations/<session_id>', methods=['PATCH'])
@jwt_required()
def update_conversation(session_id):
    """
    Update conversation title
    
    Request body:
    {
        "title": "New title"
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        session = ChatSession.query.filter_by(
            id=session_id,
            user_id=user_id
        ).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Conversation not found'
            }), 404
        
        if 'title' in data:
            session.title = data['title']
        
        session.updated_at = datetime.utcnow()
        db.session.commit()
        
        return jsonify({
            'success': True,
            'conversation': session.to_dict()
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/conversations/<session_id>', methods=['DELETE'])
@jwt_required()
def delete_conversation(session_id):
    """Delete a conversation (soft delete)"""
    try:
        user_id = get_jwt_identity()
        
        session = ChatSession.query.filter_by(
            id=session_id,
            user_id=user_id
        ).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Conversation not found'
            }), 404
        
        # Soft delete
        session.is_active = False
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Conversation deleted successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/conversations/<session_id>/messages', methods=['POST'])
@jwt_required()
def add_message_to_conversation(session_id):
    """
    Add message to conversation
    
    Request body:
    {
        "role": "user" or "assistant",
        "content": "Message text",
        "tokens_used": 150,
        "model_used": "gpt-4"
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        # Verify session belongs to user
        session = ChatSession.query.filter_by(
            id=session_id,
            user_id=user_id
        ).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Conversation not found'
            }), 404
        
        # Create message
        message = Message(
            session_id=session_id,
            role=data['role'],
            content=data['content'],
            tokens_used=data.get('tokens_used'),
            model_used=data.get('model_used')
        )
        
        db.session.add(message)
        
        # Generate title if first message
        if not session.title and data['role'] == 'user':
            session.generate_title()
        
        # Update session timestamp
        session.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': message.to_dict()
        }), 201
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


# ==================== BOOKMARKS ====================

@ux_bp.route('/bookmarks', methods=['GET'])
@jwt_required()
def get_bookmarks():
    """
    Get user's bookmarks
    
    Query params:
    - type: Filter by bookmark type (case, query, conversation, document)
    - folder: Filter by folder
    - favorites: Show only favorites (true/false)
    - search: Search in titles and notes
    """
    try:
        user_id = get_jwt_identity()
        bookmark_type = request.args.get('type')
        folder = request.args.get('folder')
        favorites_only = request.args.get('favorites', 'false').lower() == 'true'
        search_term = request.args.get('search')
        
        # Get bookmarks
        bookmarks = Bookmark.get_user_bookmarks(
            user_id=user_id,
            bookmark_type=bookmark_type,
            folder=folder,
            favorites_only=favorites_only
        )
        
        # Apply search filter if provided
        if search_term:
            search_lower = search_term.lower()
            bookmarks = [
                b for b in bookmarks
                if (b.item_title and search_lower in b.item_title.lower()) or
                   (b.notes and search_lower in b.notes.lower())
            ]
        
        # Get unique folders
        folders = db.session.query(Bookmark.folder).filter_by(
            user_id=user_id
        ).distinct().all()
        folders = [f[0] for f in folders]
        
        return jsonify({
            'success': True,
            'bookmarks': [b.to_dict() for b in bookmarks],
            'total': len(bookmarks),
            'folders': folders
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/bookmarks', methods=['POST'])
@jwt_required()
def create_bookmark():
    """
    Create a bookmark
    
    Request body:
    {
        "type": "case|query|conversation|document",
        "item_id": "unique-id",
        "title": "Bookmark title",
        "preview": "Short preview text",
        "folder": "folder-name",
        "tags": ["tag1", "tag2"],
        "notes": "User notes"
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        if not data or 'type' not in data or 'item_id' not in data:
            return jsonify({
                'error': 'Missing required fields: type, item_id'
            }), 400
        
        # Create bookmark
        bookmark = Bookmark.create_bookmark(
            user_id=user_id,
            bookmark_type=data['type'],
            item_id=data['item_id'],
            title=data.get('title'),
            preview=data.get('preview'),
            folder=data.get('folder', 'default'),
            tags=data.get('tags', []),
            notes=data.get('notes')
        )
        
        return jsonify({
            'success': True,
            'bookmark': bookmark.to_dict()
        }), 201
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


@ux_bp.route('/bookmarks/<bookmark_id>', methods=['PATCH'])
@jwt_required()
def update_bookmark(bookmark_id):
    """
    Update bookmark details
    
    Request body:
    {
        "title": "New title",
        "folder": "new-folder",
        "tags": ["tag1", "tag2"],
        "notes": "Updated notes",
        "is_favorite": true
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json()
        
        bookmark = Bookmark.query.filter_by(
            id=bookmark_id,
            user_id=user_id
        ).first()
        
        if not bookmark:
            return jsonify({
                'success': False,
                'error': 'Bookmark not found'
            }), 404
        
        # Update fields
        if 'title' in data:
            bookmark.item_title = data['title']
        if 'folder' in data:
            bookmark.folder = data['folder']
        if 'tags' in data:
            bookmark.tags = data['tags']
        if 'notes' in data:
            bookmark.notes = data['notes']
        if 'is_favorite' in data:
            bookmark.is_favorite = data['is_favorite']
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'bookmark': bookmark.to_dict()
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/bookmarks/<bookmark_id>', methods=['DELETE'])
@jwt_required()
def delete_bookmark(bookmark_id):
    """Delete a bookmark"""
    try:
        user_id = get_jwt_identity()
        
        bookmark = Bookmark.query.filter_by(
            id=bookmark_id,
            user_id=user_id
        ).first()
        
        if not bookmark:
            return jsonify({
                'success': False,
                'error': 'Bookmark not found'
            }), 404
        
        db.session.delete(bookmark)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Bookmark deleted successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@ux_bp.route('/bookmarks/<bookmark_id>/access', methods=['POST'])
@jwt_required()
def record_bookmark_access(bookmark_id):
    """Record that user accessed a bookmark"""
    try:
        user_id = get_jwt_identity()
        
        bookmark = Bookmark.query.filter_by(
            id=bookmark_id,
            user_id=user_id
        ).first()
        
        if not bookmark:
            return jsonify({
                'success': False,
                'error': 'Bookmark not found'
            }), 404
        
        bookmark.update_access()
        
        return jsonify({
            'success': True,
            'message': 'Access recorded'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== EXPORTS ====================

@ux_bp.route('/export/conversation/<session_id>', methods=['POST'])
@jwt_required()
def export_conversation(session_id):
    """
    Export conversation to PDF/DOCX/TXT/JSON
    
    Request body:
    {
        "format": "pdf|docx|txt|json",
        "include_timestamps": true,
        "include_metadata": true
    }
    """
    try:
        user_id = get_jwt_identity()
        data = request.get_json() or {}
        export_format = data.get('format', 'pdf').lower()
        
        # Verify session belongs to user
        session = ChatSession.query.filter_by(
            id=session_id,
            user_id=user_id
        ).first()
        
        if not session:
            return jsonify({
                'success': False,
                'error': 'Conversation not found'
            }), 404
        
        # Generate filename
        title_safe = "".join(c for c in (session.title or 'conversation') if c.isalnum() or c in (' ', '_', '-'))
        filename = f"conversation_{title_safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{export_format}"
        
        # Create export record
        export_record = ExportHistory.create_export(
            user_id=user_id,
            export_type='conversation',
            export_format=export_format,
            filename=filename,
            content_id=session_id,
            content_title=session.title
        )
        
        # Prepare export data
        export_data = session.to_dict(include_messages=True)
        
        try:
            # Create export based on format
            if export_format == 'json':
                output_file = _export_conversation_json(export_data, filename)
            elif export_format == 'txt':
                output_file = _export_conversation_txt(export_data, filename, data)
            elif export_format == 'pdf':
                output_file = _export_conversation_pdf(export_data, filename, data)
            elif export_format == 'docx':
                output_file = _export_conversation_docx(export_data, filename, data)
            else:
                export_record.mark_failed(f'Unsupported format: {export_format}')
                return jsonify({
                    'success': False,
                    'error': f'Unsupported format: {export_format}'
                }), 400
            
            # Mark export as completed
            file_size = os.path.getsize(output_file)
            export_record.mark_completed(file_size=file_size, file_path=output_file)
            
            # Send file
            return send_file(
                output_file,
                as_attachment=True,
                download_name=filename
            )
            
        except Exception as e:
            export_record.mark_failed(str(e))
            raise
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


@ux_bp.route('/export/history', methods=['GET'])
@jwt_required()
def get_export_history():
    """Get user's export history"""
    try:
        user_id = get_jwt_identity()
        limit = request.args.get('limit', 20, type=int)
        
        exports = ExportHistory.query.filter_by(
            user_id=user_id
        ).order_by(
            ExportHistory.created_at.desc()
        ).limit(limit).all()
        
        return jsonify({
            'success': True,
            'exports': [e.to_dict() for e in exports],
            'total': len(exports)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== HELPER FUNCTIONS ====================

def _export_conversation_json(data: dict, filename: str) -> str:
    """Export conversation as JSON"""
    exports_dir = Path('data/exports')
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = exports_dir / filename
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    return str(filepath)


def _export_conversation_txt(data: dict, filename: str, options: dict) -> str:
    """Export conversation as plain text"""
    exports_dir = Path('data/exports')
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = exports_dir / filename
    
    include_timestamps = options.get('include_timestamps', True)
    
    with open(filepath, 'w', encoding='utf-8') as f:
        # Header
        f.write("="*80 + "\n")
        f.write(f"Conversation: {data['title']}\n")
        f.write(f"Created: {data['created_at']}\n")
        f.write(f"Messages: {data['message_count']}\n")
        f.write("="*80 + "\n\n")
        
        # Messages
        for msg in data.get('messages', []):
            role = msg['role'].upper()
            content = msg['content']
            
            f.write(f"{role}:\n")
            if include_timestamps:
                f.write(f"[{msg['timestamp']}]\n")
            f.write(f"{content}\n")
            f.write("-"*80 + "\n\n")
        
        # Footer
        f.write("="*80 + "\n")
        f.write("Exported from LegalChatbot\n")
        f.write(f"Export Date: {datetime.now().isoformat()}\n")
    
    return str(filepath)


def _export_conversation_pdf(data: dict, filename: str, options: dict) -> str:
    """Export conversation as PDF (using reportlab)"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
    except ImportError:
        # Fallback to TXT if reportlab not available
        filename_txt = filename.replace('.pdf', '.txt')
        return _export_conversation_txt(data, filename_txt, options)
    
    exports_dir = Path('data/exports')
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = exports_dir / filename
    
    # Create PDF
    doc = SimpleDocTemplate(str(filepath), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        spaceAfter=30
    )
    
    # Add title
    title = Paragraph(f"Conversation: {data['title']}", title_style)
    story.append(title)
    story.append(Spacer(1, 0.2*inch))
    
    # Add metadata
    metadata_style = styles['Normal']
    story.append(Paragraph(f"<b>Created:</b> {data['created_at']}", metadata_style))
    story.append(Paragraph(f"<b>Messages:</b> {data['message_count']}", metadata_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Add messages
    for msg in data.get('messages', []):
        role = msg['role'].upper()
        content = msg['content']
        
        # Role header
        role_style = ParagraphStyle(
            'Role',
            parent=styles['Heading3'],
            textColor='blue' if role == 'USER' else 'green'
        )
        story.append(Paragraph(f"{role}:", role_style))
        
        # Content
        story.append(Paragraph(content, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
    
    # Build PDF
    doc.build(story)
    
    return str(filepath)


def _export_conversation_docx(data: dict, filename: str, options: dict) -> str:
    """Export conversation as DOCX (using python-docx)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback to TXT if python-docx not available
        filename_txt = filename.replace('.docx', '.txt')
        return _export_conversation_txt(data, filename_txt, options)
    
    exports_dir = Path('data/exports')
    exports_dir.mkdir(parents=True, exist_ok=True)
    
    filepath = exports_dir / filename
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Conversation: {data['title']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Created: {data['created_at']}")
    doc.add_paragraph(f"Messages: {data['message_count']}")
    doc.add_paragraph()
    
    # Add messages
    for msg in data.get('messages', []):
        role = msg['role'].upper()
        content = msg['content']
        
        # Role header
        role_para = doc.add_heading(f"{role}:", level=2)
        if role == 'USER':
            role_para.runs[0].font.color.rgb = RGBColor(0, 0, 255)
        else:
            role_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
        # Content
        doc.add_paragraph(content)
        doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"\nExported from LegalChatbot on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer_para.runs[0].font.size = Pt(10)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    doc.save(str(filepath))
    
    return str(filepath)


# Register blueprint
def register_ux_enhancements(app):
    """Register UX enhancements blueprint"""
    app.register_blueprint(ux_bp)
    print("✅ UX Enhancements API registered at /api/ux")
